import os
from datetime import datetime
from docx import Document
from docx.shared import Inches

def generate_report(ticket_text, ticket_info, db_findings, runbot_findings, resolution, screenshots, output_dir="output") -> str:
    try:
        os.makedirs(output_dir, exist_ok=True)
        doc = Document()
        
        doc.add_heading("Odoo Support Ticket — Investigation Report", 1)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        doc.add_heading("Ticket Summary", 2)
        doc.add_paragraph(str(ticket_info.get("summary", "")))
        doc.add_paragraph(f"Module: {ticket_info.get('module', '')}")
        doc.add_paragraph(f"Odoo Version: {ticket_info.get('odoo_version', '')}")
        doc.add_paragraph(f"Error: {ticket_info.get('error_message', '')}")
        
        doc.add_heading("Original Ticket", 2)
        doc.add_paragraph(str(ticket_text))
        
        doc.add_heading("Investigation Findings", 2)
        doc.add_paragraph(str(db_findings))
        if runbot_findings:
            doc.add_paragraph(str(runbot_findings))
            
        doc.add_heading("Screenshots", 2)
        for path in screenshots:
            if os.path.exists(path):
                doc.add_picture(path, width=Inches(5.5))
                doc.add_paragraph(path)
            else:
                doc.add_paragraph(f"[Screenshot not available: {path}]")
                
        doc.add_heading("Resolution Guide", 2)
        doc.add_paragraph(str(resolution))
        
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        
        return filepath
    except Exception as e:
        print(f"Error generating report: {e}")
        return ""
